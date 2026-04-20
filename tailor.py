import anthropic
import docx
import docx.oxml
import datetime
import os
import sys
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

RESUME = """KATHERINE PAIGE BLACK

EDUCATION
Georgetown University, Washington D.C. - MS Economics
The George Washington University, Washington D.C. - BA Economics

WORK
End to End Analytics @ Accenture - Industrial and Functional AI Decision Scientist (Aug 2021-Present)
  - Albertsons: Served as onshore lead Data Scientist for meat department markdown optimization across ~200 Intermountain West stores (hundreds of SKUs); engineered end-to-end ML pipeline integrating 4 data sources (POS transactions, forward-looking coupons, historical sales/seasonality, DC shipments) to replace a fully manual markdown process, leading a 6-person cross-functional team across two time zones.
  - Gilead Pharma: Replaced a non-functional legacy Alteryx process (requiring manual file downloads, unused by the team) with a fully automated Python/Databricks pipeline on AWS pulling dynamic cloud data; built production dashboard adopted across multiple teams measuring line loads against campaign output; introduced supplier confidence heuristic enabling data-driven contract manufacturer negotiations; delivered full handover training to client counterpart.
  - Burlington: Developed inventory optimization models using clustering and time series techniques across ~100 stores; pilot results showed 30% sales lift in select merchandise tranches vs. non-pilot stores, informing stocking decisions to streamline a bloated product offering.
  - Google BIBA: Designed supply chain forecasting dashboards for Google's full data center fleet with 30+ power users; built modular dashboard architecture that scaled with new AI infrastructure rollouts, tracking downtime and quality metrics across the entire population.

US Patent and Trademark Office - Statistical Programmer (2019-2021)
- Statistical analysis, visualizations, published report to Congress

World Bank - Consultant (Jan-May 2019)
- RCT on labor inequality in MENA

Campbell Hill Aviation Group - Research Analyst (2016-2018)
- Business cases for airports and airlines

SKILLS: Python, SQL, S&OP, machine learning, Anaplan L1MB and L2MB"""

SYSTEM_PROMPT = """You are an expert resume writer and career coach. Your job is to tailor a resume to a specific job description.

When tailoring the resume:
- The result MUST fit on one page — be ruthless about cutting less relevant details
- Keep only 2-3 bullets per job, prioritizing the most relevant to the job description
- Drop entire roles or bullets that add little value for this specific job
- Reorder and reframe bullets to highlight the most relevant experience first
- Mirror language and keywords from the job description where they authentically apply
- Keep all information truthful — do not fabricate experience or credentials
- Preserve the candidate's actual job titles, employers, and dates exactly
- Output the resume in clean plain text with NO extra blank lines between bullets
- Use a single blank line between sections, nothing more
- For the Accenture role, format project bullets with two leading spaces before the dash (e.g. "  - Albertsons: ...") to show they are sub-items under the job title"""

def get_job_description():
    """Get job description from a file arg, clipboard, or stdin pipe."""
    import subprocess

    # Option 1: file path passed as argument
    if len(sys.argv) > 1:
        path = sys.argv[1]
        with open(path, "r") as f:
            text = f.read().strip()
        if text:
            print(f"Reading job description from: {path}")
            return text

    # Option 2: piped via stdin (e.g. pbpaste | python tailor.py)
    if not sys.stdin.isatty():
        text = sys.stdin.read().strip()
        if text:
            return text

    # Option 3: read from clipboard automatically
    result = subprocess.run(["pbpaste"], capture_output=True, text=True)
    text = result.stdout.strip()
    if text:
        print("Reading job description from clipboard.")
        return text

    return ""


def read_short_input(prompt):
    """Read a short optional note from the terminal."""
    print(prompt)
    print("Press Enter to skip, or type a note and press Enter:")
    tty = open("/dev/tty", "r")
    try:
        line = tty.readline().strip()
    finally:
        tty.close()
    return line


def main():
    job_description = get_job_description()

    if not job_description:
        print("No job description found.")
        print("Usage options:")
        print("  1. Copy the job description to your clipboard, then run the script.")
        print("  2. Save it to a file and run: python tailor.py job.txt")
        print("  3. Pipe it: pbpaste | python tailor.py")
        return

    print()
    extra_notes = read_short_input("Any additional details? (e.g. projects, skills, context not on your resume)")

    print("\nTailoring your resume...\n")
    print("=" * 60)

    client = anthropic.Anthropic()

    output_text = []

    content = f"Here is my resume:\n\n{RESUME}\n\n"
    content += f"Here is the job description:\n\n{job_description}\n\n"
    if extra_notes:
        content += f"Here are some additional details I want incorporated — please weave them in naturally and format them professionally:\n\n{extra_notes}\n\n"
    content += "Please tailor my resume to this job description."

    with client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
    ) as stream:
        for block in stream:
            if (
                block.type == "content_block_delta"
                and block.delta.type == "text_delta"
            ):
                print(block.delta.text, end="", flush=True)
                output_text.append(block.delta.text)

    print("\n" + "=" * 60)

    # Save as Word document
    full_text = "".join(output_text)

    doc = docx.Document()

    # Tight page margins (1 inch → 0.75 inch)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Remove default spacing from Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(13)

    lines = full_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        # Skip consecutive blank lines — allow only one
        if line.strip() == "":
            if i > 0 and lines[i - 1].strip() != "":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Indented sub-bullet (lines starting with "  -")
        if line.startswith("  -"):
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            run = p.add_run(line.strip())

        # Bold section headers (all-caps lines like EDUCATION, WORK, SKILLS)
        elif line.isupper() and not line.startswith("-"):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(4)

        # Bold the name (first non-empty line)
        elif i == 0 or (i < 3 and all(l.strip() == "" for l in lines[:i])):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_after = Pt(4)

        else:
            run = p.add_run(line)

        i += 1

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"resume_tailored_{timestamp}.docx"
    filepath = os.path.join(os.path.expanduser("~"), "Desktop", "tailored_resumes", filename)
    doc.save(filepath)
    print(f"\nSaved to: {filepath}")

if __name__ == "__main__":
    main()
